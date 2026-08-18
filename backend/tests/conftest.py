"""Shared fixtures: generated DOCX and PDF policy documents."""

import pytest


SAMPLE_SECTIONS = [
    ("1. Introduction", 1, [
        "This policy governs data handling at Acme Corporation.",
    ]),
    ("2. Data Handling", 1, []),
    ("2.1 Retention", 2, [
        "Records must be destroyed after 90 days.",
        "Backups are retained for 30 days.",
    ]),
    ("2.2 Access Control", 2, [
        "Access to VaultMaster requires manager approval.",
        "Contact security@acme.com for exceptions.",
    ]),
]


@pytest.fixture
def sample_docx(tmp_path):
    import docx

    d = docx.Document()
    for heading, level, paras in SAMPLE_SECTIONS:
        d.add_heading(heading, level=level)
        for p in paras:
            d.add_paragraph(p)
    # bullet list under 2.2
    d.add_paragraph("Report incidents within 24 hours.", style="List Bullet")
    d.add_paragraph("Escalate breaches to the CISO.", style="List Bullet")
    # 2x2 table
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Data class"
    table.cell(0, 1).text = "Retention"
    table.cell(1, 0).text = "Financial"
    table.cell(1, 1).text = "7 years"
    path = tmp_path / "policy.docx"
    d.save(str(path))
    return path


@pytest.fixture
def sample_pdf(tmp_path):
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    y = 72
    def write(text, size, bold=False):
        nonlocal y, page
        fontname = "helvetica-bold" if bold else "helvetica"
        page.insert_text((72, y), text, fontsize=size, fontname=fontname)
        y += size * 1.6
        if y > 750:
            page = doc.new_page()
            y = 72

    write("1. Introduction", 18, bold=True)
    write("This policy governs data handling at Acme Corporation.", 11)
    write("2. Data Handling", 18, bold=True)
    write("2.1 Retention", 14, bold=True)
    write("Records must be destroyed after 90 days.", 11)
    write("Backups are retained for 30 days.", 11)
    write("2.2 Access Control", 14, bold=True)
    write("Access requires manager approval.", 11)
    path = tmp_path / "policy.pdf"
    doc.save(str(path))
    doc.close()
    return path
